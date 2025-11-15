import io
import json
import base64  
from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.opc.exceptions import OpcError
from docx.text.paragraph import Paragraph

app = Flask(__name__)

def extract_skills_from_doc(doc: Document):
    existing_skills = set()
    skills_heading_found = False

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                skills_heading_found_in_cell = False
                for i, para in enumerate(cell.paragraphs):
                    current_text = para.text.strip()
                    if not current_text:
                        continue
                    if not skills_heading_found_in_cell:
                        if current_text.lower().startswith('skills'):
                            skills_heading_found_in_cell = True
                        continue
                    if ',' in current_text:
                        skills = current_text.split(',')
                        for skill in skills:
                            cleaned_skill = skill.strip().title()
                            if cleaned_skill:
                                existing_skills.add(cleaned_skill)
                        break
                    else:
                        is_new_heading = (para.style.name.startswith('Heading') or 
                                          (current_text.isupper() and len(current_text) > 4))
                        if is_new_heading:
                            break
                        
                        cleaned_skill = current_text.strip().title()
                        if cleaned_skill:
                            existing_skills.add(cleaned_skill)
                
                if skills_heading_found_in_cell: break
            if skills_heading_found_in_cell: break
        if skills_heading_found_in_cell: break

    if not existing_skills:
        skills_heading_found = False
        for i, para in enumerate(doc.paragraphs):
            current_text = para.text.strip()
            if not current_text:
                continue
            if not skills_heading_found:
                if current_text.lower().startswith('skills') and not para.style.name.lower().startswith('list'):
                    skills_heading_found = True
                continue
            if ',' in current_text:
                skills = current_text.split(',')
                for skill in skills:
                    cleaned_skill = skill.strip().title()
                    if cleaned_skill:
                        existing_skills.add(cleaned_skill)
                break
        
            else:
                is_new_heading = (para.style.name.startswith('Heading') or 
                                  (current_text.isupper() and len(current_text) > 4))
                if is_new_heading:
                    break
                
                cleaned_skill = current_text.strip().title()
                if cleaned_skill:
                    existing_skills.add(cleaned_skill)

    print(f"Found existing skills: {list(existing_skills)}")
    return list(existing_skills)

def get_processed_skills(skills_to_add):
    processed_skills = []
    for skill_string in skills_to_add:
        individual_skills = skill_string.split(',')
        for skill in individual_skills:
            cleaned_skill = skill.strip().title() 
            if cleaned_skill:
                processed_skills.append(cleaned_skill)
    return processed_skills

def _copy_paragraph_format(source_para: Paragraph, target_para: Paragraph):
    fmt = source_para.paragraph_format
    target_fmt = target_para.paragraph_format
    target_fmt.alignment = fmt.alignment
    target_fmt.first_line_indent = fmt.first_line_indent
    target_fmt.keep_together = fmt.keep_together
    target_fmt.keep_with_next = fmt.keep_with_next
    target_fmt.left_indent = fmt.left_indent
    target_fmt.line_spacing = fmt.line_spacing
    target_fmt.line_spacing_rule = fmt.line_spacing_rule
    target_fmt.page_break_before = fmt.page_break_before
    target_fmt.right_indent = fmt.right_indent
    target_fmt.space_after = fmt.space_after
    target_fmt.space_before = fmt.space_before
    target_fmt.widow_control = fmt.widow_control


def add_skills_to_paragraph(para: Paragraph, skills_to_add_list: list):
    new_skills_string = ", ".join(skills_to_add_list)
    existing_text = para.text.strip()
    
    if existing_text and not existing_text.endswith(','):
        text_to_add = f", {new_skills_string}"
    else:
        text_to_add = f" {new_skills_string}"
    
    new_run = para.add_run(text_to_add)
    
    if para.runs and len(para.runs) > 1:
        last_run = para.runs[-2] 
        font = last_run.font
        new_run.font.name = font.name
        new_run.font.size = font.size
        new_run.font.bold = font.bold
        new_run.font.italic = font.italic
        new_run.font.underline = font.underline
        new_run.font.color.rgb = font.color.rgb


def add_skills_to_list(insertion_para: Paragraph, skills_to_add_list: list, ref_para: Paragraph):
    for skill in reversed(skills_to_add_list):
        try:
            new_para = insertion_para.insert_paragraph_before(skill, style=ref_para.style)
            _copy_paragraph_format(ref_para, new_para)
        except ValueError:
            new_para = insertion_para.insert_paragraph_before(skill)
            _copy_paragraph_format(ref_para, new_para)

@app.route("/add-skills", methods=["POST"])
def modify_resume():
    
    if 'resume' not in request.files:
        return jsonify({"error": "No resume file provided"}), 400
        
    file_storage = request.files['resume']
    skills_to_add_raw = request.form.getlist('skills') 
    if not skills_to_add_raw:
        return jsonify({"error": "No skills provided"}), 400

    try:
        file_bytes = file_storage.stream.read()
        doc_for_reading = Document(io.BytesIO(file_bytes))
        
        doc_for_writing = Document(io.BytesIO(file_bytes))
        
        existing_skills = extract_skills_from_doc(doc_for_reading)

        processed_skills_to_add = get_processed_skills(skills_to_add_raw)

        existing_skills_lower = set(s.lower() for s in existing_skills)
        final_skills_to_add = [
            s for s in processed_skills_to_add 
            if s.lower() not in existing_skills_lower
        ]

        if not final_skills_to_add:
            print("No new skills to add. All recommended skills are already present.")
        
        skills_added = False
        
        # Table Search Mode
        for table in doc_for_writing.tables:
            for row in table.rows:
                for cell in row.cells:
                    skills_heading_found_in_cell = False
                    first_skill_para_index = -1 
                    
                    for i, para in enumerate(cell.paragraphs):
                        current_text = para.text.lower().strip()
                        if not skills_heading_found_in_cell:
                            if current_text.startswith('skills'):
                                skills_heading_found_in_cell = True
                            continue 
                        if current_text:
                            first_skill_para_index = i
                            break
                    
                    if first_skill_para_index != -1:
                        first_skill_para = cell.paragraphs[first_skill_para_index]
                        is_paragraph_style = ',' in first_skill_para.text
                        
                        if not is_paragraph_style:
                            # Add to list in cell
                            insertion_point = None
                            reference_para = first_skill_para 
                            
                            for j in range(first_skill_para_index + 1, len(cell.paragraphs)):
                                sub_para = cell.paragraphs[j]
                                sub_text = sub_para.text.strip()
                                
                                is_blank = not sub_text
                                is_new_heading = (sub_para.style.name.startswith('Heading') or 
                                                  (sub_text.isupper() and sub_text and len(sub_text) > 4)) 
                                
                                if is_blank or is_new_heading:
                                    insertion_point = sub_para
                                    if j > first_skill_para_index:
                                        reference_para = cell.paragraphs[j-1]
                                    break
                                elif sub_text:
                                    reference_para = sub_para 
                            
                            if insertion_point:
                                add_skills_to_list(insertion_point, final_skills_to_add, reference_para)
                            else:
                                for skill in final_skills_to_add:
                                    new_para = cell.add_paragraph(skill, style=reference_para.style)
                                    _copy_paragraph_format(reference_para, new_para)
                            skills_added = True

                        else:
                            add_skills_to_paragraph(first_skill_para, final_skills_to_add)
                            skills_added = True
                    
                    if skills_added: break
                if skills_added: break
            if skills_added: break
        
        # Paragraph Search Mode
        if not skills_added:
            skills_heading_found = False
            target_para_index = -1 
            
            for i, para in enumerate(doc_for_writing.paragraphs):
                current_text = para.text.lower().strip()
                if not skills_heading_found:
                    if current_text.startswith('skills') and not para.style.name.lower().startswith('list'):
                        skills_heading_found = True
                    continue
                if current_text: 
                    target_para_index = i
                    break
            
            if target_para_index != -1:
                target_para = doc_for_writing.paragraphs[target_para_index]
                is_paragraph_style = ',' in target_para.text
                
                if not is_paragraph_style:
                    # Add to list in body
                    insertion_point = None
                    reference_para = target_para 
                    
                    for j in range(target_para_index + 1, len(doc_for_writing.paragraphs)):
                        sub_para = doc_for_writing.paragraphs[j]
                        sub_text = sub_para.text.strip()
                        is_blank = not sub_text
                        is_new_heading = (sub_para.style.name.startswith('Heading') or 
                                          (sub_text.isupper() and sub_text and len(sub_text) > 4)) 
                        
                        if is_blank or is_new_heading:
                            insertion_point = sub_para
                            if j > target_para_index:
                                reference_para = doc_for_writing.paragraphs[j-1]
                            break
                        elif sub_text:
                            reference_para = sub_para
                    
                    if insertion_point:
                        add_skills_to_list(insertion_point, final_skills_to_add, reference_para)
                    else: 
                        for skill in final_skills_to_add:
                            new_para = doc_for_writing.add_paragraph(skill, style=reference_para.style)
                            _copy_paragraph_format(reference_para, new_para)
                    skills_added = True
                
                else:
                    add_skills_to_paragraph(target_para, final_skills_to_add)
                    skills_added = True

        # Fallback
        if not skills_added and final_skills_to_add:
            doc_for_writing.add_heading('Skills', level=1)
            new_skills_string = ", ".join(final_skills_to_add)
            doc_for_writing.add_paragraph(new_skills_string)
            print("Fallback: Added skills to end of document.")

        file_stream = io.BytesIO()
        doc_for_writing.save(file_stream)
        file_stream.seek(0)
        
        encoded_file_data = base64.b64encode(file_stream.read()).decode('utf-8')
        
        print("Successfully processed resume. Returning JSON.")
        
        return jsonify({
            "existing_skills": existing_skills,
            "file_data": encoded_file_data,
            "added_skills": final_skills_to_add
        })
        
    except OpcError:
        return jsonify({"error": "Invalid file format. Please upload a .docx file."}), 400
    except Exception as e:
        print(f"An error occurred: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

if __name__ == "__main__":
    app.run(debug=True, port=5000)